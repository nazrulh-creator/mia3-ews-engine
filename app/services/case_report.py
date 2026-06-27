"""One-page Early-Warning Case Report (.docx) for a flagged account.

Adapted from MicroFlex's Credit Decision Report: a deterministic, on-engine
narrative by default; server-rendered charts via Pillow (no headless browser).
An optional Claude-written summary is OFF by default and blocked on LIVE, so
real borrower data never leaves the engine.
"""
from __future__ import annotations

import io
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from app.config import get_settings
from app.core.scoring import BAND_META
from app.db.models import AccountScore


def _band_color(band: str) -> RGBColor:
    hexv = BAND_META.get(band, {}).get("color", "#666666").lstrip("#")
    return RGBColor(int(hexv[0:2], 16), int(hexv[2:4], 16), int(hexv[4:6], 16))


def _factor_chart(score: AccountScore) -> Optional[io.BytesIO]:
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None
    factors = score.top_factors or []
    if not factors:
        return None
    W, H, pad, rowh = 520, 26 * len(factors) + 20, 10, 22
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    maxc = max((abs(float(f["contribution"])) for f in factors), default=1.0) or 1.0
    mid = W // 2
    for i, f in enumerate(factors):
        y = pad + i * rowh
        c = float(f["contribution"])
        length = int((abs(c) / maxc) * (W // 2 - 80))
        color = (192, 57, 43) if c > 0 else (39, 174, 96)
        if c > 0:
            d.rectangle([mid, y, mid + length, y + rowh - 8], fill=color)
        else:
            d.rectangle([mid - length, y, mid, y + rowh - 8], fill=color)
        d.text((4, y), str(f["label"])[:34], fill="black")
    d.line([mid, 0, mid, H], fill=(150, 150, 150))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def build_case_report(score: AccountScore, *, run_ref: str) -> bytes:
    settings = get_settings()
    doc = Document()

    title = doc.add_heading("Early-Warning Case Report", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"{settings.owner_team} · {settings.app_name}").italic = True
    if not settings.is_live:
        warn = doc.add_paragraph()
        r = warn.add_run("TEST ENVIRONMENT — NOT LIVE DATA")
        r.bold = True
        r.font.color.rgb = RGBColor(0xB5, 0x61, 0x0C)  # CGC Orange (readable on white)

    meta = doc.add_paragraph()
    meta.add_run(f"Account {score.account_id}").bold = True
    meta.add_run(f"   ·   FI {score.fi_id} ({score.fi_name or '—'})"
                 f"   ·   {score.scheme} / {score.sector}"
                 f"   ·   Run {run_ref}")

    band_p = doc.add_paragraph()
    br = band_p.add_run(f"Classification: {score.band}")
    br.bold = True
    br.font.size = Pt(14)
    br.font.color.rgb = _band_color(score.band)
    band_p.add_run(f"    Risk score {score.risk_score:.2f}    "
                   f"Confidence {score.confidence:.0f}/100 ({score.confidence_band})")

    # Risk arithmetic — the live worked example.
    doc.add_heading("How the band was reached", level=2)
    bd = score.breakdown or {}
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Component", "Value", "Rank", "Weighted"]):
        hdr[i].text = h
    rows = [
        ("Probability of MIA3 slip", f"{score.probability*100:.1f}%", score.pd_rank,
         f"{bd.get('w_pd', 0.5)}×{score.pd_rank} = {bd.get('pd_term', 0):.2f}"),
        ("Exposure at Default", f"RM {score.ead:,.0f}", score.ead_rank,
         f"{bd.get('w_ead', 0.3)}×{score.ead_rank} = {bd.get('ead_term', 0):.2f}"),
        ("Outstanding Ratio", f"{score.outstanding_ratio*100:.1f}%", score.outratio_rank,
         f"{bd.get('w_outratio', 0.2)}×{score.outratio_rank} = {bd.get('out_term', 0):.2f}"),
    ]
    for name, val, rank, weighted in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = name, val, str(rank), weighted
    total = table.add_row().cells
    total[0].text, total[3].text = "Risk score", f"{score.risk_score:.2f}"

    doc.add_paragraph(
        f"Score {score.risk_score:.2f} crosses the cut-off for "
        f"“{score.band}”. {BAND_META.get(score.band, {}).get('definition', '')}")

    # Why the model scored it.
    doc.add_heading("Why the model flagged this account", level=2)
    doc.add_paragraph(score.explanation_operational or "")
    chart = _factor_chart(score)
    if chart is not None:
        doc.add_picture(chart, width=Inches(5.2))

    doc.add_heading("Recommended handling", level=2)
    handling = {
        "needs_review": "Hold for mandatory human review (borderline confidence).",
        "fast_track": "High-confidence high-risk — action via the branch worklist.",
        "no_review": "No review required at current confidence.",
        "reviewed": "A reviewer decision has been recorded for this account.",
    }.get(score.review_status, "Review per standard procedure.")
    doc.add_paragraph(handling)

    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Generated by the MIA3 Early-Warning Engine. A flag is a prompt for "
        "human attention, not an automated decision. Model is in calibration; "
        "precision is intentionally traded for recall.")
    fr.italic = True
    fr.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
