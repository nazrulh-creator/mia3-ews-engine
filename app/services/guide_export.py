"""Render the in-app User Guide and Quick Start to downloadable .docx.

Single source of truth: the guide is generated from app.guide_content.SECTIONS
(the same content the web pages use), so the document never drifts from the
app. Inline SVG diagrams become a captioned line (the rich, interactive
versions live in the in-app guide); everything else — headings, prose, step
lists, tables, callouts — is rendered faithfully.
"""
from __future__ import annotations

import io
from html.parser import HTMLParser
from typing import List, Optional

from docx import Document
from docx.shared import Pt, RGBColor

from app import guide_content

SDA = RGBColor(0x2B, 0x3A, 0x55)
ACCENT = RGBColor(0x46, 0x61, 0x8F)
_INLINE = {"b", "strong", "em", "i", "code", "span", "a"}


def _clean(text: str) -> str:
    return " ".join(text.split())


class _HtmlToDocx(HTMLParser):
    """Convert the limited HTML used in guide bodies into docx blocks."""

    def __init__(self, doc: Document):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.in_svg = False
        self.mode: Optional[str] = None      # p | note | warn | seenh | li | h2 | h3
        self.buf = ""
        self.list_kind: Optional[str] = None  # ul | ol
        self.list_items: List[str] = []
        self.in_table = False
        self.rows: List[List[str]] = []
        self.row: List[str] = []
        self.in_cell = False
        self.cell = ""

    # -- helpers --
    def _flush_block(self):
        text = _clean(self.buf)
        if text:
            if self.mode == "h2":
                self.doc.add_heading(text, level=2)
            elif self.mode == "h3":
                self.doc.add_heading(text, level=3)
            elif self.mode == "seenh":
                par = self.doc.add_paragraph(); run = par.add_run(text); run.bold = True
            elif self.mode in ("note", "warn"):
                prefix = "Note — " if self.mode == "note" else "Caution — "
                par = self.doc.add_paragraph(); run = par.add_run(prefix + text)
                run.italic = True; run.font.color.rgb = ACCENT
            elif self.mode == "p":
                self.doc.add_paragraph(text)
        self.buf = ""
        self.mode = None

    # -- tag handlers --
    def handle_starttag(self, tag, attrs):
        if self.in_svg:
            return
        ad = dict(attrs)
        cls = ad.get("class", "")
        if tag == "svg":
            self.in_svg = True
            label = ad.get("aria-label", "Diagram")
            par = self.doc.add_paragraph(); r = par.add_run(f"[ Diagram: {label} — see the in-app guide ]")
            r.italic = True; r.font.color.rgb = ACCENT; r.font.size = Pt(9.5)
            return
        if tag in _INLINE:
            return
        if tag == "br":
            self.buf += " "
            return
        if tag in ("h2", "h3"):
            self._flush_block(); self.mode = tag
        elif tag == "p":
            self._flush_block(); self.mode = "seenh" if "g-seen-h" in cls else "p"
        elif tag == "div":
            if "g-note" in cls:
                self._flush_block(); self.mode = "note"
            elif "g-warn" in cls:
                self._flush_block(); self.mode = "warn"
        elif tag in ("ul", "ol"):
            self.list_kind = tag; self.list_items = []
        elif tag == "li":
            self.mode = "li"; self.buf = ""
        elif tag == "table":
            self.in_table = True; self.rows = []
        elif tag == "tr":
            self.row = []
        elif tag in ("th", "td"):
            self.in_cell = True; self.cell = ""

    def handle_data(self, data):
        if self.in_svg:
            return
        if self.in_cell:
            self.cell += data
        elif self.mode is not None:
            self.buf += data

    def handle_endtag(self, tag):
        if tag == "svg":
            self.in_svg = False
            return
        if self.in_svg or tag in _INLINE or tag == "br":
            return
        if tag in ("h2", "h3", "p"):
            self._flush_block()
        elif tag == "div":
            if self.mode in ("note", "warn"):
                self._flush_block()
        elif tag == "li":
            text = _clean(self.buf)
            if text:
                self.list_items.append(text)
            self.buf = ""; self.mode = None
        elif tag in ("ul", "ol"):
            style = "List Number" if self.list_kind == "ol" else "List Bullet"
            for item in self.list_items:
                self.doc.add_paragraph(item, style=style)
            self.list_kind = None; self.list_items = []
        elif tag in ("th", "td"):
            self.row.append(_clean(self.cell)); self.in_cell = False
        elif tag == "tr":
            if self.row:
                self.rows.append(self.row)
        elif tag == "table":
            self._render_table()
            self.in_table = False

    def _render_table(self):
        if not self.rows:
            return
        ncol = max(len(r) for r in self.rows)
        table = self.doc.add_table(rows=0, cols=ncol)
        table.style = "Light Grid Accent 1"
        for ri, row in enumerate(self.rows):
            cells = table.add_row().cells
            for ci in range(ncol):
                val = row[ci] if ci < len(row) else ""
                cells[ci].text = ""
                run = cells[ci].paragraphs[0].add_run(val)
                if ri == 0:
                    run.bold = True
        self.doc.add_paragraph()


def _title_block(doc: Document, title: str, subtitle: str):
    h = doc.add_paragraph(); r = h.add_run(title)
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = SDA
    s = doc.add_paragraph(); rs = s.add_run(subtitle); rs.italic = True


def build_guide_docx() -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    _title_block(doc, "MIA3 Early Warning Engine — User Guide",
                 "Strategic Data Analytics (SDA) · a team within CGC Malaysia")
    doc.add_paragraph(
        "A step-by-step walkthrough of every screen in the engine. The interactive "
        "version with diagrams is in the app at /guide; this document is a printable "
        "companion.")
    for s in guide_content.SECTIONS:
        doc.add_heading(f"{s['num']}. {s['title']}", level=1)
        parser = _HtmlToDocx(doc)
        parser.feed(s["body"])
        parser.close()
    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def build_quickstart_docx() -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    _title_block(doc, "MIA3 Early Warning Engine — Quick Start",
                 "Smoke-detector for the guaranteed portfolio · human-on-the-loop · fully auditable")
    panels = [
        ("1 · Open it & sign in", [
            "Hosted: mia3-ews-engine-test.fly.dev (first load wakes the server, ~3–5s).",
            "Demo logins (TEST, synthetic): internal/internal123, checker/checker123, "
            "branch/branch123, fi_mbb/fi123."]),
        ("2 · Environments — check the colour", [
            "LIVE = green bar (real monitoring). TEST = amber bar + watermark + TEST- ids "
            "(synthetic sandbox).",
            "Separate databases; test data never reaches LIVE."]),
        ("3 · Your first look (30 seconds)", [
            "Sign in as internal.",
            "Demo → Generate & score demo portfolio.",
            "Open the Dashboard — read the band cards and the by-FI / by-sector tables.",
            "Click a total → open an account → read its explanation and arithmetic."]),
        ("4 · What MIA3 does", [
            "Watches loans already on the book and scores monthly which are likely to slip "
            "into Months-in-Arrears 3, then ranks them by potential impact."]),
        ("5 · The flow", [
            "Data arrives → validated against the contract → model scores P(MIA3) → risk "
            "score 0.5/0.3/0.2 → four bands → confidence routing → three role-based views.",
            "Bands: Low <2.0 · Moderate 2–3 · High 3–3.5 · Very High >3.5."]),
        ("6 · Models, ensembles & governed settings", [
            "Models: register synthetic, glass-box logistic/OLS (enter coefficients) or "
            "uploaded ML models — one or more active per segment.",
            "Decision rules: when several models are active, a governed rule (average, "
            "weighted, max, majority…) combines them into the trigger.",
            "Weights, band cut-offs, calibration, models and rules are all dual-controlled; "
            "the dashboard shows which models and rules are live."]),
        ("7 · Safety rails", [
            "Precision is very low (~0.1) by design — high recall means every high-risk flag "
            "goes to a human, never an automated action.",
            "Borderline-confidence cases are held for review.",
            "Every step is written to an append-only, hash-chained audit trail."]),
        ("8 · Get help", [
            "Click ‘User guide for this page’ in any screen's banner for that screen's section.",
            "Open the Guide link in the top nav any time; hover the i icons on form fields.",
            "Toggle guided / compact in the header to show or hide help."]),
    ]
    for title, items in panels:
        doc.add_heading(title, level=2)
        for it in items:
            doc.add_paragraph(it, style="List Bullet")
    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()
