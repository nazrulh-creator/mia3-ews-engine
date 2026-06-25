"""Generate the MIA3 Early Warning Engine — Software Design Document.

Single source of content rendered to two outputs:
  * a polished Word .docx (cover, TOC field, page numbers) → ~/Downloads
  * a version-controlled Markdown copy → docs/SOFTWARE_DESIGN_DOCUMENT.md

    python -m scripts.make_sdd [output_docx_dir]

The content is authored from the implemented system, so it is ground-truth.
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

VERSION = "1.0"
ENGINE_VERSION = "0.9.0"
DATE = "June 2026"

# --- Content model: a flat list of blocks -------------------------------- #
# ("h", level, text) | ("p", text) | ("b", [items]) | ("n", [items])
# ("t", [headers], [[cells]]) | ("code", text) | ("note", text)

C = []
def h(level, text): C.append(("h", level, text))
def p(text): C.append(("p", text))
def b(items): C.append(("b", items))
def n(items): C.append(("n", items))
def t(headers, rows): C.append(("t", headers, rows))
def code(text): C.append(("code", text))
def note(text): C.append(("note", text))

# ========================================================================= #
h(1, "1. Document control")
t(["Field", "Value"], [
    ["Title", "MIA3 Early Warning Engine — Software Design Document"],
    ["Version", VERSION + f" (engine v{ENGINE_VERSION})"],
    ["Status", "Draft for review"],
    ["Owner", "Strategic Data Analytics (SDA), a team within CGC Malaysia"],
    ["Prepared for", "Nazrul — Chief Strategy Officer"],
    ["Audience", "Credit-risk officers, model owners, validators, approvers, engineers, auditors"],
    ["Companion documents", "MIA3 Build & Deployment Plan; MIA3 — Adopting & Adapting the "
     "MicroFlex Patterns; EWS deck (MIA3 Predictor Model & Risk Classification Framework)"],
    ["Repository", "github.com/nazrulh-creator/mia3-ews-engine"],
    ["Hosted (TEST)", "mia3-ews-engine-test.fly.dev"],
])
note("This document describes a system that productionises a back-tested model. "
     "It is written to be defensible to a board and a regulator: every governed "
     "behaviour, safeguard and audit mechanism is specified, not assumed.")

h(1, "2. Introduction")
h(2, "2.1 Purpose and scope")
p("This Software Design Document (SDD) specifies the design of the MIA3 Early Warning "
  "Engine — a standalone portfolio-monitoring platform that turns the back-tested MIA3 "
  "predictor model into a working tool for CGC's financial-institution (FI) partners, "
  "branches and internal risk team. It covers the architecture, data design, functional "
  "subsystems, governance mechanisms, security, interfaces, operations and quality "
  "assurance. It does not cover the statistical development of the predictor model "
  "itself, which is documented in the EWS deck and the model's back-testing record.")
h(2, "2.2 Relationship to MicroFlex")
p("MIA3 deliberately reuses the proven governance and engineering patterns of the imSME "
  "MicroFlex Decisioning & Routing Engine, but answers a different question. The single "
  "most useful distinction:")
b(["MicroFlex is a gatekeeper — it decides whether to let a new loan through the door "
   "(eligibility, pricing, routing), reacting to one applicant in real time.",
   "MIA3 is a smoke detector — it watches the loans already on the book and raises an "
   "alarm when one starts to smoulder, sweeping the whole portfolio on a schedule."])
p("Consequently MIA3 adopts MicroFlex's engine-agnostic governance (dual control, "
  "hash-chained audit, golden test pack, LIVE/TEST split, four-layer help) almost "
  "verbatim, adapts its confidence, calibration and explainability ideas to a "
  "monitoring context, and leaves behind its decisioning-only features (eligibility "
  "gating, guarantee pricing, applicant routing, bureau ensembles).")
h(2, "2.3 Design principles")
n(["Governance first — every consequential change is versioned, reasoned, dual-controlled, "
   "and recorded on a tamper-evident trail.",
   "Human-on-the-loop — while precision is low, every high-risk flag is a prompt for human "
   "attention, never an automated action against a customer.",
   "Settings over code — band cut-offs, component weights and calibration are governed "
   "settings, never edited in source.",
   "Modular monolith — one well-organised, deployable application with clean internal "
   "boundaries; split into services only if load ever demands it.",
   "No real borrower data in development or external AI tools — synthetic or de-identified "
   "data only.",
   "Explain as you go — a four-layer help system and a context-sensitive guide assume no "
   "prior training."])
h(2, "2.4 The precision caution (a first-class design driver)")
p("The model is tuned to catch most genuine cases (recall ~0.73) while keeping false "
  "negatives low — a sound choice for an early-warning tool. The trade-off is very low "
  "precision (~0.1 in recent monthly results): roughly nine in ten flagged accounts will "
  "be false alarms. This is acceptable "
  "for a tool that prioritises accounts for a human to look at, and unacceptable for one "
  "that takes automated action. The entire design keeps a person in the loop on every "
  "high-risk call while calibration continues; the confidence-based review flow (§6.4) is "
  "the direct, structural answer to this caution.")

h(1, "3. System overview")
h(2, "3.1 Objectives")
n(["Predict the probability that a guaranteed account slips into Months-in-Arrears 3 "
   "(MIA 3) within the forecast horizon.",
   "Classify each account into one of four risk bands by combining that probability with "
   "exposure and leverage, so attention goes to the accounts that matter most."])
h(2, "3.2 Users and views")
t(["Role", "Sees", "Can do"], [
    ["Internal Risk", "The whole portfolio", "Everything, plus governed tuning, model "
     "registry, runs and audit"],
    ["Branch", "The whole portfolio (action worklist)", "Review decisions; no tuning"],
    ["Financial Institution (FI)", "Only its own book (row-scoped by fi_id)", "Read its "
     "accounts and explanations; JSON API"],
])
h(2, "3.3 High-level capability map")
b(["Monthly batch scoring of the guaranteed portfolio (file / DB feed / manual / scheduled).",
   "Risk-score classification into four bands per the EWS framework.",
   "Per-account and per-decision explainability (SHAP + the live risk arithmetic).",
   "Confidence-based routing to a human review queue.",
   "Governed tuning, calibration and model lifecycle, all dual-controlled and audited.",
   "Portfolio early-warning ladder, learnings library, printable case reports.",
   "Three role-based views, a self-documenting JSON API, and a context-sensitive user guide."])

h(1, "4. Architecture")
h(2, "4.1 Architectural style")
p("MIA3 is a modular monolith. It separates three concerns internally — the scoring path "
  "(high-volume, automated), the configuration/governance path (low-volume, dual-controlled) "
  "and the human path (review) — but ships as one deployable application with clean module "
  "boundaries. This mirrors MicroFlex's advice to start as a modular monolith and split "
  "later only if load demands it.")
h(2, "4.2 Layered component model")
t(["Layer", "Modules", "Responsibility"], [
    ["Scoring core", "app/core/ — features, model, scoring, validation, confidence, "
     "explain, batch, synthetic", "Pure logic: contract, model loading, risk score, "
     "validation, confidence, explanation. No web or DB dependency."],
    ["Persistence", "app/db/ — database, models, audit", "ORM, the hash-chained audit "
     "store, sessions."],
    ["Services", "app/services/ — governance, runs, seed, case_report", "Run orchestration, "
     "dual-control governance, persistence of scores, bootstrap, report generation."],
    ["Web", "app/routers/ + app/templates/", "Role-based views, the JSON API, the in-app "
     "user guide."],
    ["Entrypoints", "scripts/ — score_file, run_batch, build_golden_pack, "
     "train_synthetic_model, make_test_portfolios, make_sdd", "CLI scorer, scheduled run, "
     "tooling."],
])
h(2, "4.3 The five coupling rules")
p("These boundaries are adopted verbatim from MicroFlex; each prevents a specific way an "
  "audit can go wrong.")
n(["The scoring engine reads only the approved, active model version — never a draft.",
   "The explainability layer reads the stored decision record, never a recomputation "
   "against today's model — so an explanation always reflects the score as it was made.",
   "The audit store is append-only and kept separate from everyday data updates.",
   "A human override changes the treatment, never the model's number — the estimate and "
   "the decision are recorded as two distinct facts.",
   "Direct database edits to models or settings are locked out; everything flows through "
   "the governed workflow."])
h(2, "4.4 End-to-end data flow")
n(["Data arrives — by default a monthly file (CSV or JSON) is uploaded; a live DB feed, a "
   "manual run, or a scheduled run are alternatives.",
   "The engine validates the data against the contract; malformed rows are quarantined, "
   "never guessed.",
   "The active model scores every accepted account's probability of MIA 3 slip; SHAP "
   "contributions are computed and stored in the same pass (off the critical path).",
   "The risk score is computed — 0.5·rank(P) + 0.3·rank(EAD) + 0.2·rank(Outstanding ratio) "
   "— and the account is banded.",
   "A confidence score and a routing decision are attached.",
   "Results are persisted, timestamped and tagged with the model version; portfolio "
   "early-warning ladder alerts are raised.",
   "The three role-based views read from the same persisted numbers."])
h(2, "4.5 Technology stack and rationale")
t(["Choice", "Why"], [
    ["Python", "The model is a Python (XGBoost) object; serving it from Python avoids a "
     "translation layer. Natural fit for a Claude Code build."],
    ["FastAPI", "Fast, self-documenting web layer (OpenAPI at /api/docs) — useful when an "
     "FI integrates."],
    ["SQLAlchemy + PostgreSQL", "Scores, audit and settings live in a proper database so "
     "they can be sliced by FI, scheme and sector. SQLite is the zero-install local/TEST "
     "default via the same ORM."],
    ["Fly.io", "Same platform as MicroFlex; the web app runs as one machine and scheduled "
     "scoring as a machine that wakes, works and sleeps."],
    ["python-docx + Pillow", "Server-rendered one-page case reports with charts, no "
     "headless browser."],
    ["XGBoost / SHAP / LIME", "Production model serving and explanation; the engine "
     "degrades gracefully to a deterministic stand-in and exact contributions when these "
     "libraries are absent, so it is always demonstrable."],
])

h(1, "5. Data design")
h(2, "5.1 The data contract")
p("The data contract (app/core/features.py) is the single source of truth for the monthly "
  "input file. It defines three column groups and a missing-value policy. The "
  "machine-readable contract is exposed at /api/v1/contract; the human-readable version at "
  "/contract.")
t(["Group", "Columns"], [
    ["Identity & roll-up", "account_id, fi_id, fi_name, scheme, sector, branch, as_of_date"],
    ["Risk-score inputs", "ead (Exposure at Default, RM), outstanding_ratio (0–1)"],
    ["Model features (12)", "mia, prev_delinquency_count, arrears_movement_trend, "
     "payment_consistency, cost_to_income_ratio, profit_margin, length_of_business_months, "
     "outstanding_amount, utilization_ratio, debt_pressure_to_remaining, "
     "repayment_stress_ratio, payment_gap_x_pd"],
])
p("Missing-value policy, per column:")
t(["Policy", "Behaviour"], [
    ["reject", "The row cannot be scored safely and is quarantined."],
    ["default", "The value is filled with the documented default and the account's "
     "confidence is reduced."],
])
note("Replacing the synthetic model with the real artifact means aligning MODEL_FEATURES "
     "with the artifact's feature list; the loader checks the two agree and refuses to "
     "score on a mismatch. The full data dictionary is in Appendix A.")
h(2, "5.2 Persistence model")
p("Eleven tables, created on first run; additive column migrations are applied at startup "
  "(SQLite/Postgres ALTER TABLE ADD COLUMN for new columns).")
t(["Table", "Purpose"], [
    ["users", "Accounts, roles, FI scope, UI mode."],
    ["model_registry", "Registered model versions, status, back-test metrics, artifact "
     "path, registrant/approver."],
    ["threshold_config", "Versioned 50/30/20 weights and band cut-offs; dual-controlled."],
    ["calibration_config", "Versioned calibration mapping (identity/linear/platt); "
     "dual-controlled."],
    ["scoring_runs", "One row per run: source, model, counts, quality report, checkpoint "
     "status."],
    ["account_scores", "One row per scored account: probability, ranks, risk score, band, "
     "confidence, routing, stored explanation and features."],
    ["review_decisions", "Human review outcomes (confirm/override/escalate) with reason."],
    ["audit_events", "Append-only, hash-chained event log."],
    ["learnings", "Outcomes, reviewer notes and change reasoning."],
    ["portfolio_alerts", "Early-warning ladder alerts at FI/sector level."],
    ["problem_reports", "Universal error reports with audit linkage."],
])
h(2, "5.3 Environment-separated data")
p("LIVE and TEST use separate databases and separate audit chains (distinct "
  "MIA3_DATABASE_URL). Test data can never reach LIVE. The default environment is TEST.")

h(1, "6. Functional design")
h(2, "6.1 Ingestion and validation")
p("Uploaded files are read (CSV or JSON) and validated against the contract before any "
  "scoring. Validation produces a set of accepted rows, a set of quarantined rows, and a "
  "data-quality report (rows in/accepted/quarantined/defaulted, acceptance rate, issues "
  "by type). A structurally invalid file (missing a hard-required column) is quarantined "
  "whole.")
h(2, "6.2 Model abstraction and loading")
p("The engine never talks to a concrete model object directly; it talks to a ModelWrapper. "
  "Two implementations exist: a RealModel that wraps the back-tested artifact (XGBoost "
  "Booster, scikit-learn API, or pickle) pointed to by the active registry entry or "
  "MIA3_MODEL_PATH; and a deterministic SyntheticModel stand-in with transparent, "
  "auditable logistic coefficients that needs no ML libraries and gives reproducible "
  "probabilities (so the golden pack is stable). Loading order: active registry artifact → "
  "env artifact → synthetic stand-in. The loader validates that the model's features match "
  "the data contract.")
h(2, "6.3 Scoring core — risk score and bands")
p("The risk score combines three ranked components (each ranked 1–4, low→high risk):")
code("Risk score = 0.50 · rank(P(MIA3 slip)) + 0.30 · rank(EAD) + 0.20 · rank(Outstanding ratio)")
t(["Rank", "P(MIA3 slip)", "EAD (RM)", "Outstanding ratio"], [
    ["1", "< 25%", "< 50k", "< 25%"],
    ["2", "25–50%", "50k–200k", "25–50%"],
    ["3", "50–75%", "200k–500k", "50–75%"],
    ["4", "> 75%", "> 500k", "> 75%"],
])
t(["Band", "Score", "Definition"], [
    ["Very High Risk", "> 3.5", "Very likely to cause significant loss; immediate action."],
    ["High Risk", "3.0 – 3.5", "Likely to cause loss with noticeable impact."],
    ["Moderate Risk", "2.0 – 3.0", "Early signs of risk; still manageable."],
    ["Low Risk", "< 2.0", "Financially stable."],
])
note("Weights and cut-offs are loaded from the active threshold_config, not hardcoded. The "
     "deck values are the defaults.")
h(2, "6.4 Confidence model and routing")
p("Every prediction carries a 0–100 confidence score — a weighted blend of five components: "
  "35% model performance, 25% data completeness, 20% data quality, 10% population fit, 10% "
  "calibration. Confidence routes work: high-confidence high-risk fast-tracks to the "
  "worklist; borderline confidence is held for mandatory review; low-risk high-confidence "
  "needs no review. Thresholds: high ≥ 70, low < 55.")
h(2, "6.5 Explainability")
b(["SHAP is the official explanation — computed for every scored account during the batch "
   "run and stored, so dashboards open instantly and explanations are reproducible.",
   "LIME is an on-demand diagnostic challenger for a single account.",
   "Each explanation renders at three audience levels — technical (raw features, for "
   "validators), operational (plain phrases, for branch officers) and simplified.",
   "When SHAP is unavailable, the engine falls back to the synthetic model's exact "
   "log-odds contributions, so an explanation is always available."])
h(2, "6.6 Batch scoring engine and run lifecycle")
p("score_frame() drives validation → scoring → risk → confidence → explanation with no "
  "web or DB dependency, so it is unit-testable and reused by the CLI, the scheduled run "
  "and the web upload. execute_run() wraps it with persistence, audit and the ladder. A "
  "run is held for sign-off (a checkpoint) if requested, if nothing could be scored, or if "
  "acceptance falls below 75%; otherwise it publishes to the views.")
h(2, "6.7 Early-warning case report")
p("A one-page Word (.docx) case report is generated on demand for any account: the live "
  "risk arithmetic, the SHAP factors with a server-rendered contribution chart, the "
  "confidence, and the recommended handling. The narrative is deterministic and "
  "on-engine; an optional Claude-written summary is OFF by default and force-disabled on "
  "LIVE, so real borrower data never leaves the engine.")
h(2, "6.8 Portfolio early-warning ladder")
p("Beyond per-account flags, the engine watches concentrations. When the share of an FI's "
  "or a sector's book in the high-risk bands crosses a governed level it raises a "
  "portfolio-level alert: Watch at ≥ 15%, Halt at ≥ 30% (groups smaller than five accounts "
  "are ignored).")

h(1, "7. Governance design")
h(2, "7.1 Model registry and lifecycle")
p("Model versions move through states draft → active → retired. Registration (maker) "
  "creates a draft with an artifact (uploaded or path) and back-test metrics. Activation "
  "(checker) is dual-controlled — the registrant cannot activate their own model — and "
  "pre-flights the artifact against the data contract before swapping it in; the previous "
  "active model is retired and the next run uses the new one. Retiring the live model is "
  "single-control and immediate (the safe direction); scoring falls back to the synthetic "
  "stand-in. The active model is locked against edits until retired.")
h(2, "7.2 Threshold and weight tuning")
p("The 50/30/20 weights and band cut-offs are changed through a maker→preview→checker flow. "
  "A proposal must pass validation (weights sum to 1.0; cut-offs ordered), can be previewed "
  "against the latest portfolio (exact re-banding using stored ranks — no re-scoring), and "
  "is activated by a different approver. Every step is versioned and audited.")
h(2, "7.3 Calibration layer")
p("A governed calibration layer maps raw model output onto observed reality. Methods: "
  "identity (none, the default so nothing changes silently), linear (a·p + b) and Platt "
  "(scale/shift on the log-odds). Calibration is proposed and approved under dual control "
  "and recorded; it is the right home for back-testing findings that the model runs hot or "
  "cold.")
h(2, "7.4 Hash-chained audit store")
p("Every consequential action is appended as an event whose SHA-256 hash covers its own "
  "content plus the previous event's hash, per environment. Any later edit or deletion "
  "breaks the chain and is detectable by verify_chain(), which recomputes the chain and "
  "reports the first break. Timestamps are stored naive-UTC so they round-trip identically "
  "through the database. The log is append-only and read-only from the interface, and "
  "exportable for an auditor.")
code("hash = SHA256( seq | ts | env | actor | action | entity | before | after | detail | prev_hash )")
h(2, "7.5 Dual control")
p("Nothing consequential happens on one person's say-so. Threshold, calibration and model "
  "activation each require a second, different approver; self-approval is blocked. Changes "
  "are versioned, reasoned, time-stamped and reversible, and the safe direction (turning "
  "something back off) is always available immediately.")
h(2, "7.6 Safe degradation")
p("A case that cannot be scored correctly, or cannot be recorded, is never decided "
  "automatically. Each foreseeable failure has a defined response:")
t(["Failure", "Defined response"], [
    ["File malformed / unreadable", "Reject the run; surface the error; nothing persisted."],
    ["Required column absent", "Whole file structurally quarantined; missing columns recorded."],
    ["Required value missing in a row", "That row quarantined; the rest score normally."],
    ["Defaultable value missing", "Filled with the documented default; confidence reduced."],
    ["Acceptance rate < 75%", "Run held for sign-off (checkpoint), not auto-published."],
    ["Borderline confidence on an elevated band", "Routed to mandatory human review."],
    ["SHAP / XGBoost libraries absent", "Fall back to exact contributions / synthetic model."],
    ["Audit store unreachable", "Scoring transaction rolls back — no score without a record."],
])

h(1, "8. Security and access control")
h(2, "8.1 Authentication")
p("Session-based authentication (signed cookies, 8-hour expiry). Passwords are hashed with "
  "PBKDF2-SHA256 (pure-Python, no native backend). Credential entry is the user's own; the "
  "system never stores plaintext credentials.")
h(2, "8.2 Roles and row-level scoping")
p("Three roles — internal, branch, fi. FI users are row-scoped to their own fi_id "
  "everywhere (dashboard, accounts, review, API). Internal-risk is the only role that can "
  "reach tuning, the model registry, runs and audit.")
h(2, "8.3 Environment ringfencing")
p("LIVE shows a green bar; TEST shows an amber bar, a watermark and TEST- run ids, with a "
  "separate database and audit chain. Anything that is not exactly LIVE is treated as TEST, "
  "so LIVE is never reached by accident.")
h(2, "8.4 Data protection")
p("No real borrower data is used in development, testing or external AI tools — synthetic "
  "or de-identified data only. The optional LLM narrative is OFF by default and blocked on "
  "LIVE. Nothing destructive or outbound (LIVE deploy, dispatch to FIs) happens without "
  "explicit human go-ahead.")

h(1, "9. User-experience design")
h(2, "9.1 Four-layer help system")
p("Guidance is delivered in four layers on every screen: a plain-language purpose banner, "
  "an info tooltip on every data-entry field, guided/compact modes (a per-user toggle), and "
  "a context-sensitive user guide. A build-time check (a test) fails if any registered "
  "screen lacks a purpose or any declared data-entry field lacks a tooltip.")
h(2, "9.2 In-app user guide and quick start")
p("An 18-section user guide, with SVG visual cues, has a stable anchor per screen. Every "
  "screen's purpose banner carries a contextual link that deep-links to the section "
  "explaining that screen; a persistent Guide link and a Quick Start card are always "
  "reachable.")
h(2, "9.3 Universal problem reporting and branding")
p("Every screen has a problem-report control that captures the screen, user and time and "
  "links to the audit trail. The application carries SDA team ownership marks, not CGC "
  "corporate branding, identifying it as an SDA-built internal tool.")

h(1, "10. Interfaces")
h(2, "10.1 Web routes (selected)")
t(["Route", "Purpose"], [
    ["/dashboard", "Portfolio overview, breakdowns, trend, ladder alerts"],
    ["/accounts, /accounts/{id}", "Account list/worklist; account detail + case report"],
    ["/review", "Human review queue"],
    ["/runs", "Upload / manual run; run history + data-quality report"],
    ["/tuning", "Threshold + calibration dual-control workflow"],
    ["/models, /models/*", "Model registry: register, edit, activate, retire"],
    ["/audit", "Hash-verified audit trail"],
    ["/learnings, /demo, /contract", "Learnings library; demonstration mode; data contract"],
    ["/guide, /quickstart", "Context-sensitive user guide and quick start"],
])
h(2, "10.2 JSON API")
p("A self-documenting API (OpenAPI at /api/docs), row-scoped for FI callers: "
  "/api/v1/contract (the canonical schema), /api/v1/portfolio/summary, /api/v1/accounts, "
  "/api/v1/accounts/{account_id}.")
h(2, "10.3 Input-file interface")
p("CSV or JSON matching the data contract. JSON accepts a top-level array or a "
  "{\"accounts\": [...]} envelope. A file's SHA-256 fingerprint is recorded on the run.")

h(1, "11. Operational design")
h(2, "11.1 Deployment")
p("Containerised (python:3.11-slim) and deployed to Fly.io as a separate app with its own "
  "database. The web process runs uvicorn; the monthly scoring runs as a scheduled machine "
  "(scripts.run_batch) or via the Runs page. The TEST app uses SQLite on a mounted volume; "
  "production uses managed PostgreSQL shared by the web and scheduled processes.")
h(2, "11.2 Configuration")
t(["Variable", "Meaning"], [
    ["MIA3_ENV", "LIVE or TEST (default TEST)."],
    ["MIA3_DATABASE_URL", "Database connection; SQLite default locally."],
    ["MIA3_SECRET_KEY", "Session signing secret."],
    ["MIA3_MODEL_PATH", "Optional path to a model artifact (overrides the synthetic stand-in)."],
    ["MIA3_ENABLE_LLM_NARRATIVE", "Optional Claude narrative; OFF by default, blocked on LIVE."],
])
h(2, "11.3 Migrations and storage")
p("Schema is created on startup; additive column migrations run idempotently. Uploaded "
  "model artifacts are stored under the data volume so they persist across restarts.")

h(1, "12. Quality assurance")
h(2, "12.1 The golden test pack")
p("Because the code is AI-generated, it is treated as untrusted until proven. A golden pack "
  "of fully-specified accounts with frozen expected probability, risk score, band and "
  "confidence is reproduced exactly on every change; any divergence blocks the release. The "
  "synthetic model's determinism makes the pack stable.")
h(2, "12.2 Automated tests")
b(["test_scoring — the deck's worked example, band boundaries, rank edges, weight validation.",
   "test_pipeline — end-to-end scoring, quarantine of malformed rows, borderline routing.",
   "test_golden — exact reproduction of the golden pack.",
   "test_audit_chain — hash-chain verification and tamper detection.",
   "test_help_coverage — every screen has a purpose; every data-entry field has a tooltip; "
   "every screen maps to a guide section."])

h(1, "13. Non-functional requirements")
t(["Attribute", "Design response"], [
    ["Auditability", "Hash-chained append-only log; stored explanations; versioned governed settings."],
    ["Maintainability", "Modular monolith; pure scoring core; reused MicroFlex conventions."],
    ["Portability", "Same code on SQLite (local) and Postgres (prod); graceful ML-library fallback."],
    ["Scalability", "Batch scoring is vectorised; scheduled machine scales independently of the web app."],
    ["Reliability", "Safe-degradation table; checkpoints; no score without a record."],
    ["Usability", "Four-layer help; guided/compact; contextual guide; role-appropriate views."],
])

h(1, "14. Risks, limitations and assumptions")
b(["Very low precision (~0.1) — roughly nine in ten flags are false alarms; mitigated "
   "structurally by human-in-the-loop and the confidence gate.",
   "The hosted TEST app runs a synthetic stand-in model; results are illustrative, not real "
   "monitoring, until the back-tested artifact is dropped in.",
   "Calibration is pending back-testing outcomes; the calibration layer exists but defaults "
   "to uncalibrated.",
   "The TEST volume is single (not redundant); production should use managed Postgres and "
   "durable artifact storage.",
   "Expected-loss economics are intentionally out of scope at launch (a possible later "
   "addition)."])

h(1, "15. Build sequence (as delivered)")
t(["Phase", "Scope", "Status"], [
    ["1", "Scoring core (contract, model, risk score, validation, confidence, explain)", "Delivered"],
    ["2", "Database and hash-chained audit", "Delivered"],
    ["3", "Web layer and the three role-based views", "Delivered"],
    ["4", "Explainability and decision transparency", "Delivered"],
    ["5", "Review flow and dual-control tuning", "Delivered"],
    ["6", "Demo mode, learnings, scheduling, deployment", "Delivered"],
    ["+", "In-app user guide, tooltips, model-registry management", "Delivered"],
])

h(1, "16. Roadmap")
b(["Promote the real back-tested XGBoost artifact via the registry / MIA3_MODEL_PATH.",
   "Managed Postgres in production so the scheduled scorer shares state with the web app.",
   "Model-drift watch — a monthly check on input drift against the training distribution.",
   "Scheduled per-FI exports; a live database feed connector.",
   "Optional expected-loss economics on at-risk exposure."])

h(1, "Appendix A — Data dictionary")
t(["Column", "Type", "Required", "On missing", "Notes"], [
    ["account_id", "str", "yes", "reject", "Unique account / application number"],
    ["fi_id", "str", "yes", "reject", "FI code; row-level access key"],
    ["fi_name", "str", "no", "default", "FI display name"],
    ["scheme", "str", "yes", "reject", "Guarantee scheme code"],
    ["sector", "str", "yes", "reject", "Economic sector"],
    ["branch", "str", "no", "default", "Originating CGC branch"],
    ["as_of_date", "str", "no", "default", "Snapshot date (YYYY-MM-DD)"],
    ["ead", "float", "yes", "reject", "Exposure at Default (RM), ≥ 0"],
    ["outstanding_ratio", "float", "yes", "reject", "Outstanding / utilisation (0–1.5)"],
    ["mia", "int", "yes", "reject", "Months in arrears at snapshot (0–2)"],
    ["prev_delinquency_count", "int", "yes", "default 0", "Prior delinquency episodes (12m)"],
    ["arrears_movement_trend", "int", "yes", "default 0", "-1 improving / 0 flat / 1 worsening"],
    ["payment_consistency", "float", "yes", "default 0.5", "On-time payment share (0–1)"],
    ["cost_to_income_ratio", "float", "yes", "default 0.7", "Cost-to-income ratio"],
    ["profit_margin", "float", "yes", "default 0.05", "Net profit margin (can be negative)"],
    ["length_of_business_months", "int", "yes", "default 36", "Business age in months"],
    ["outstanding_amount", "float", "yes", "reject", "Current outstanding financing (RM)"],
    ["utilization_ratio", "float", "yes", "default 0.5", "Drawn / approved limit (0–2)"],
    ["debt_pressure_to_remaining", "float", "yes", "default 0.5", "Debt-service pressure"],
    ["repayment_stress_ratio", "float", "yes", "default 0.5", "Repayment obligation vs capacity"],
    ["payment_gap_x_pd", "float", "yes", "default 0.1", "Payment-gap × baseline-PD interaction"],
])

h(1, "Appendix B — Glossary")
t(["Term", "Meaning"], [
    ["MIA 3", "Months-in-Arrears 3 — the deterioration event the model predicts."],
    ["P(MIA3 slip)", "Model probability an account reaches MIA 3 in the horizon."],
    ["EAD", "Exposure at Default — RM amount at risk if the account defaults."],
    ["Outstanding ratio", "Outstanding / utilisation — a customer-leverage signal."],
    ["Risk score", "0.5·rank P + 0.3·rank EAD + 0.2·rank Outstanding ratio (1.0–4.0)."],
    ["Confidence", "0–100 trust in a prediction; five-component blend."],
    ["Calibration", "Governed mapping of raw model output onto observed reality."],
    ["Quarantine", "A row that cannot be scored safely; set aside, never guessed."],
    ["Golden test pack", "Frozen, signed-off expected results reproduced on every change."],
    ["Dual control", "A second, different approver required for a governed change."],
])


# --- Renderers ----------------------------------------------------------- #
def render_markdown(path: Path) -> None:
    out = [f"# MIA3 Early Warning Engine — Software Design Document\n",
           f"*Version {VERSION} · engine v{ENGINE_VERSION} · {DATE} · "
           f"Strategic Data Analytics (SDA)*\n"]
    for blk in C:
        kind = blk[0]
        if kind == "h":
            out.append(("#" * (blk[1] + 1)) + " " + blk[2] + "\n")
        elif kind == "p":
            out.append(blk[1] + "\n")
        elif kind == "b":
            out.append("\n".join(f"- {i}" for i in blk[1]) + "\n")
        elif kind == "n":
            out.append("\n".join(f"{i+1}. {v}" for i, v in enumerate(blk[1])) + "\n")
        elif kind == "code":
            out.append("```\n" + blk[1] + "\n```\n")
        elif kind == "note":
            out.append("> **Note —** " + blk[1] + "\n")
        elif kind == "t":
            headers, rows = blk[1], blk[2]
            out.append("| " + " | ".join(headers) + " |")
            out.append("| " + " | ".join("---" for _ in headers) + " |")
            for r in rows:
                out.append("| " + " | ".join(str(c) for c in r) + " |")
            out.append("")
    path.write_text("\n".join(out))


def render_docx(path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    SDA = RGBColor(0x2B, 0x3A, 0x55)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    # Cover page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        doc.add_paragraph()
    r = title.add_run("MIA3 Early Warning Engine")
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = SDA
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Software Design Document"); rs.font.size = Pt(18); rs.font.color.rgb = SDA
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {VERSION} · engine v{ENGINE_VERSION} · {DATE}\n"
                 "Strategic Data Analytics (SDA) · a team within CGC Malaysia").italic = True
    doc.add_page_break()

    # TOC field (Word populates on open / right-click → Update Field)
    doc.add_heading("Contents", level=1)
    para = doc.add_paragraph()
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t"); fld_text.text = "Right-click and choose “Update Field” to build the table of contents."
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, fld_text, fld_end):
        run._r.append(el)
    doc.add_page_break()

    # Footer page numbers
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("MIA3 EWS Engine — Software Design Document   ·   Page ")
    fr = footer.add_run()
    b1 = OxmlElement("w:fldChar"); b1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    e1 = OxmlElement("w:fldChar"); e1.set(qn("w:fldCharType"), "end")
    fr._r.append(b1); fr._r.append(it); fr._r.append(e1)

    # Body
    for blk in C:
        kind = blk[0]
        if kind == "h":
            doc.add_heading(blk[2], level=blk[1])
        elif kind == "p":
            doc.add_paragraph(blk[1])
        elif kind == "b":
            for i in blk[1]:
                doc.add_paragraph(i, style="List Bullet")
        elif kind == "n":
            for i in blk[1]:
                doc.add_paragraph(i, style="List Number")
        elif kind == "code":
            cp = doc.add_paragraph(); cr = cp.add_run(blk[1])
            cr.font.name = "Consolas"; cr.font.size = Pt(9.5)
        elif kind == "note":
            np_ = doc.add_paragraph(); nr = np_.add_run("Note — " + blk[1])
            nr.italic = True; nr.font.color.rgb = RGBColor(0x46, 0x61, 0x8F)
        elif kind == "t":
            headers, rows = blk[1], blk[2]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            for i, hdr in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(hdr); run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
            doc.add_paragraph()
    doc.save(str(path))


def main(argv) -> int:
    out_dir = Path(argv[0]) if argv else Path(os.path.expanduser("~/Downloads"))
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "MIA3_Software_Design_Document.docx"
    md_path = Path(__file__).resolve().parent.parent / "docs" / "SOFTWARE_DESIGN_DOCUMENT.md"
    render_docx(docx_path)
    render_markdown(md_path)
    print(f"Wrote SDD ({len([b for b in C if b[0]=='h'])} headings):")
    print(f"  Word : {docx_path}")
    print(f"  Markdown : {md_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
